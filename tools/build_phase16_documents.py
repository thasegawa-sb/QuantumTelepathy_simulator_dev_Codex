"""Build the two Japanese Phase 16 Word deliverables from validated artifacts."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "deliverables" / "phase16"
MANUAL_PATH = OUTPUT_DIR / "quantum_telepathy_simulator_manual_ja.docx"
REPORT_PATH = OUTPUT_DIR / "quantum_telepathy_research_report_ja.docx"
PAPER_PATH = OUTPUT_DIR / "operational_lctc_quantum_advantage.tex"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x20, 0x37, 0x48)
MUTED = RGBColor(0x66, 0x6F, 0x78)
GOLD = RGBColor(0xA6, 0x78, 0x20)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def load_json(relative_path: str) -> dict[str, Any]:
    return json.loads((ROOT / relative_path).read_text(encoding="utf-8"))


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def set_run_font(
    run: Any,
    *,
    name: str = "Noto Sans CJK JP",
    east_asia: str = "Noto Sans CJK JP",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell: Any, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, *, top: int, bottom: int, start: int, end: int) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table: Any, widths_inches: Sequence[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), "9360")
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(round(width * 1440)))
        grid.append(grid_column)
    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        row_properties.append(no_split)
        for cell, width in zip(row.cells, widths_inches, strict=True):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(round(width * 1440)))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell, top=80, bottom=80, start=120, end=120)


def add_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_instruction = OxmlElement("w:instrText")
    field_instruction.set(qn("xml:space"), "preserve")
    field_instruction.text = instruction
    field_separator = OxmlElement("w:fldChar")
    field_separator.set(qn("w:fldCharType"), "separate")
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    for node in (field_begin, field_instruction, field_separator, field_text, field_end):
        run._r.append(node)
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc: Document, *, title: str, subject: str, justified: bool) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans CJK JP"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK JP")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK JP")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK JP")
    normal.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justified else WD_ALIGN_PARAGRAPH.LEFT
    )
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8 if justified else 6)
    normal.paragraph_format.line_spacing = 1.333 if justified else 1.25

    heading_settings = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12 if justified else 14, 6 if justified else 7),
        "Heading 3": (12, DARK_BLUE, 8 if justified else 10, 4 if justified else 5),
    }
    for style_name, (size, color, before, after) in heading_settings.items():
        style = doc.styles[style_name]
        style.font.name = "Noto Sans CJK JP"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK JP")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK JP")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK JP")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Noto Sans CJK JP"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK JP")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194 if justified else -0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208 if justified else 1.25

    settings = doc.settings._element
    language = settings.find(qn("w:themeFontLang"))
    if language is None:
        language = OxmlElement("w:themeFontLang")
        settings.append(language)
    language.set(qn("w:val"), "ja-JP")
    language.set(qn("w:eastAsia"), "ja-JP")

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    left = header_paragraph.add_run(title)
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = header_paragraph.add_run("\tVersion 1.0 | 2026-09-03")
    set_run_font(right, size=8.5, color=MUTED)
    tabs = header_paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.45))
    paragraph_properties = header_paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "B7C9D6")
    borders.append(bottom)
    paragraph_properties.append(borders)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_run = footer_paragraph.add_run("Quantum Network Simulator R&D Project | ")
    set_run_font(footer_run, size=8.5, color=MUTED)
    add_field(footer_paragraph, "PAGE")
    separator = footer_paragraph.add_run(" / ")
    set_run_font(separator, size=8.5, color=MUTED)
    add_field(footer_paragraph, "NUMPAGES")

    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = "Quantum Network Simulator R&D Project"
    doc.core_properties.keywords = "LCTC, CHSH, HFT, quantum network, Ding-Jiang, Li et al."
    doc.core_properties.comments = "Generated from Phase 15 validated artifacts."


def add_cover(
    doc: Document,
    *,
    kicker: str,
    title: str,
    subtitle_lines: Sequence[str],
    audience: str,
) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(96)
    kicker_paragraph = doc.add_paragraph()
    kicker_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_paragraph.paragraph_format.space_after = Pt(18)
    run = kicker_paragraph.add_run(kicker)
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(10)
    run = title_paragraph.add_run(title)
    set_run_font(run, size=28, color=INK, bold=True)

    for line in subtitle_lines:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(3)
        run = subtitle.add_run(line)
        set_run_font(run, size=14, color=DARK_BLUE)

    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_before = Pt(24)
    descriptor.paragraph_format.space_after = Pt(72)
    run = descriptor.add_run(audience)
    set_run_font(run, size=10.5, color=MUTED, italic=True)

    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.paragraph_format.space_after = Pt(4)
    run = date_paragraph.add_run("2026年9月3日 | Version 1.0")
    set_run_font(run, size=11, color=INK, bold=True)
    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project.add_run("Quantum Network Simulator R&D Project")
    set_run_font(run, size=9.5, color=MUTED)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> Any:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> Any:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True)
        body = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(body)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_code(doc: Document, lines: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_together = True
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EFF3F6")
    paragraph._p.get_or_add_pPr().append(shading)
    run = paragraph.add_run(lines)
    set_run_font(
        run,
        name="Courier New",
        east_asia="Noto Sans CJK JP",
        size=8.5,
        color=INK,
    )


def add_status_callout(doc: Document, text: str, *, fill: str = "E8F2EA") -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK, bold=True)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_inches: Sequence[float],
    *,
    header_fill: str,
    font_size: float = 8.5,
) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, color=INK, bold=True)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    set_table_width(table, widths_inches)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc: Document, relative_path: str, caption: str, alt_text: str) -> None:
    path = ROOT / relative_path
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    picture = run.add_picture(str(path), width=Inches(6.05))
    picture._inline.docPr.set("descr", alt_text)
    picture._inline.docPr.set("title", caption)
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(2)
    caption_paragraph.paragraph_format.space_after = Pt(9)
    caption_paragraph.paragraph_format.keep_with_next = False
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=8.5, color=MUTED, italic=True)


def add_contents(doc: Document, sections: Sequence[str]) -> None:
    add_heading(doc, "目次", 1)
    add_body(doc, "本版の目次は章構成を固定表示しています。Wordの見出しナビゲーションから各章へ移動できます。")
    for section in sections:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(section)
        set_run_font(run, color=INK)


def build_manual() -> None:
    phase15 = load_json(
        "experiments/final_validation/results/phase15_v1/final_validation_summary.json"
    )
    table3 = load_json("experiments/li2026/results/table3_50km_v1/table3_summary.json")
    fig3 = load_json("experiments/ding_jiang/results/fig3_v3/fig3_summary.json")
    test_count = next(
        token
        for token in phase15["validations"]["test_suite"]["summary"].split()
        if token.isdigit()
    )
    doc = Document()
    configure_document(
        doc,
        title="量子テレパシー・シミュレータ 操作マニュアル",
        subject="Ding-Jiang v3 / Li et al. v1 reproduction simulator manual",
        justified=False,
    )
    add_cover(
        doc,
        kicker="SIMULATOR MANUAL",
        title="量子テレパシー・シミュレータ",
        subtitle_lines=("操作・再現・検証マニュアル", "Ding–Jiang v3 / Li et al. v1"),
        audience="研究者・検証担当者・研究ソフトウェア開発者向け",
    )
    add_status_callout(
        doc,
        f"検証状態: Phase 15 PASS | 再現ジョブ 12/12 PASS | "
        f"既存テスト {test_count} 件 PASS",
    )
    add_body(
        doc,
        "本書は、理論上の非局所ゲーム利得から、忠実度・有限統計・局所判断遅延・エンタングルメント供給率を同時に満たす運用上の量子優位までを評価するための利用手順をまとめたものです。数値結果の意味、再現コマンド、設定変更、判定の読み方、再検証方法を対象とします。",
    )
    add_table(
        doc,
        ("項目", "固定内容"),
        (
            ("対象論文", "Ding and Jiang, arXiv:2407.21723v3"),
            ("運用モデル", "Li et al., arXiv:2604.07451v1"),
            ("Python", ">= 3.11"),
            ("最終検証日", "2026-09-03"),
            ("成果物ハッシュ", phase15["committed_artifact_manifest"]["aggregate_sha256"][:20] + "…"),
        ),
        (1.45, 5.05),
        header_fill="E8EEF5",
    )
    add_contents(
        doc,
        (
            "目的と判定境界",
            "セットアップとクイックスタート",
            "リポジトリ構成",
            "科学モデルと時間尺度",
            "設定ファイル",
            "Ding–Jiang再現手順",
            "Li et al.再現手順",
            "運用上の量子優位判定",
            "ハードウェア・M2イベント駆動モデル",
            "結果の読み方",
            "検証・回帰・拡張",
            "トラブルシューティング",
        ),
    )

    add_heading(doc, "1. 目的と判定境界", 1)
    add_body(
        doc,
        "このシミュレータは、入力分布 P(x,y)、効用 u(a,b|x,y)、局所古典戦略、量子相関、物理誤差、有限統計、ネットワーク供給能力を分離して扱います。単に Δω>0 を得ただけでは「運用上の量子優位」と報告しません。",
    )
    add_bullets(
        doc,
        (
            "theoretical_advantage: 物理誤差を含めた量子効用が古典最適値を上回るか。",
            "fidelity_criterion: 結合インフィデリティ ε が ε_th=1-C/Q 未満か。",
            "rate_criterion: R_HEG が n_req/T_env を厳密に上回るか。",
            "decision_criterion: τ_dec=τ_rot+τ_meas が T_loc 未満か。",
            "latency_constrained_regime: T_loc<T_comm が成立し、観測後通信なしの前提が有効か。",
            "overall_operational_quantum_advantage: 必要な全条件がPASSしたときのみPASS。",
        ),
    )
    add_status_callout(
        doc,
        "重要: PASSは実装済み・設定済みの範囲に対する判定です。論文点データが未公開の図や、微視的デバイス情報が不足する項目のPARTIALを上書きしません。",
        fill="FFF4D8",
    )

    add_heading(doc, "2. セットアップとクイックスタート", 1)
    add_heading(doc, "2.1 前提環境", 2)
    add_bullets(
        doc,
        (
            "Python 3.11以上。Phase 15記録環境はPython 3.13.5 / macOS 15.5 arm64。",
            "必須依存: NumPy、SciPy、CVXPY。図生成にはMatplotlib。",
            "実験スクリプトはリポジトリルートから実行する。",
        ),
    )
    add_code(
        doc,
        "python3 -m venv .venv\n"
        "source .venv/bin/activate\n"
        "python3 -m pip install -e '.[plots]'\n"
        "python3 -m pytest",
    )
    add_heading(doc, "2.2 最短の科学確認", 2)
    add_numbers(
        doc,
        (
            "CHSH単体テストを実行し、古典値0.75、量子値0.8535533906、ギャップ0.1035533906を確認する。",
            "Ding–Jiang Figure 3を再生成し、最大ギャップと古典決定論戦略列挙の一致を確認する。",
            "Li Figure 2、Figure 3、Table IIIの順に実行し、理論・統計・ハードウェアの各層を確認する。",
            "最後にPhase 15ランナーを実行し、全回帰をまとめて監査する。",
        ),
    )
    add_code(
        doc,
        "python3 -m pytest tests/scientific/test_chsh_oracles.py -q\n"
        "PYTHONPATH=src python3 experiments/ding_jiang/reproduce_fig3.py\n"
        "PYTHONPATH=src python3 -m experiments.li2026.reproduce_fig2\n"
        "python3 experiments/final_validation/run_phase15.py",
    )

    add_heading(doc, "3. リポジトリ構成", 1)
    add_table(
        doc,
        ("パス", "役割"),
        (
            ("src/quantum_telepathy/core", "一般非局所ゲーム、XORゲーム、古典戦略列挙"),
            ("src/quantum_telepathy/ding_jiang", "HFT効用、損失、v3 Type IIメモリ、ノイズ"),
            ("src/quantum_telepathy/li2026", "一般化LCTC、忠実度、統計、運用判定、三者拡張"),
            ("src/quantum_telepathy/hardware", "M0/M1/M2、HEG、Ybシステム、イベント駆動"),
            ("src/quantum_telepathy/optimization", "有限格子ハードウェア探索とPareto解析"),
            ("experiments/*/configs", "追跡可能な入力パラメータ"),
            ("experiments/*/oracles", "論文値・解析値・許容差"),
            ("experiments/*/results", "JSON、CSV、PNGの再現成果物"),
            ("docs/research", "モデル対応、仮定、検証境界、再現行列"),
            ("tests/scientific", "独立オラクルと回帰テスト"),
        ),
        (2.25, 4.25),
        header_fill="E8EEF5",
        font_size=8.2,
    )

    add_heading(doc, "4. 科学モデルと時間尺度", 1)
    add_heading(doc, "4.1 層構造", 2)
    add_numbers(
        doc,
        (
            "Layer 0: P(x,y)、u(a,b|x,y)、P(a,b|x,y)と期待効用。",
            "Layer 1: 二値入出力XORゲーム、行列M、C(M)、Q(M)、Δω。",
            "Layer 2: Ding–Jiangの理想化HFT/LCTCモデル。",
            "Layer 3: Liの非対称効用、相関入力、T_loc、T_comm、T_env。",
            "Layer 4: ε_s、ε_meas、結合ε、メモリデコヒーレンス。",
            "運用層: 有限統計、HEG率、判断遅延、M2メモリ制約。",
        ),
    )
    add_heading(doc, "4.2 混同してはいけない時間", 2)
    add_table(
        doc,
        ("記号", "意味", "判定"),
        (
            ("T_loc", "アプリケーションが許す局所判断窓", "τ_dec<T_loc"),
            ("T_comm", "拠点間通信時間", "T_loc<T_comm"),
            ("T_env", "市場・環境が定常とみなせる窓", "n_req/T_env"),
            ("τ_dec", "基底回転と測定の合計", "τ_rot+τ_meas"),
            ("τ_occ", "M2メモリ占有時間", "供給率・必要メモリ数"),
            ("τ_mem", "量子メモリ寿命", "デコヒーレンス後ε<ε_th"),
        ),
        (0.8, 3.75, 1.95),
        header_fill="E8EEF5",
    )

    add_heading(doc, "5. 設定ファイル", 1)
    add_body(
        doc,
        "全実験パラメータはJSON設定に記録します。既存設定を変更する代わりに、新しい識別子と出力先を持つ設定を追加してください。論文値はoraclesへ、シミュレータ入力はconfigsへ分離します。",
    )
    add_table(
        doc,
        ("分類", "代表パラメータ", "単位・制約"),
        (
            ("ゲーム/市場", "P(x,y), beta1, beta2, utility", "確率和=1、betaは設定範囲内"),
            ("LCTC時間", "T_loc, T_comm, T_env", "秒、T_loc<T_comm"),
            ("統計", "alpha, n_req, R_req", "0<alpha<1、R_req=n_req/T_env"),
            ("量子誤差", "epsilon_s, epsilon_meas, epsilon", "Liの厳密結合式を使用"),
            ("ネットワーク", "distance, loss, N_channels", "km、dB/km、正整数"),
            ("デバイス", "eta_det, eta_opt, tau_mem", "効率0–1、寿命は秒"),
        ),
        (1.2, 2.7, 2.6),
        header_fill="E8EEF5",
    )
    add_heading(doc, "5.1 新規シナリオの作り方", 2)
    add_numbers(
        doc,
        (
            "最も近いconfigs JSONを複製し、reference/versionとprovenanceを保持する。",
            "変更したパラメータと単位を明示し、論文値を入力へ逆算して合わせない。",
            "別のoutputディレクトリへ実行し、既存回帰成果物を上書きしない。",
            "解析解、独立列挙、Decimal直計算、公開値の順でオラクルを追加する。",
            "差異はFAIL/PARTIAL/INSUFFICIENT_INFORMATIONのまま記録する。",
        ),
    )

    add_heading(doc, "6. Ding–Jiang再現手順", 1)
    add_table(
        doc,
        ("対象", "コマンド", "主オラクル"),
        (
            ("Fig. 3", "reproduce_fig3.py", "Theorem 10、16戦略列挙、CHSH"),
            ("Sec. 4.1", "reproduce_loss_example.py", "Eq. A.11/A.12"),
            ("Fig. 5", "reproduce_fig5_cross_sections.py", "明示量子戦略、Q1+AB上界"),
            ("Type II", "reproduce_type_ii_memory.py", "v3二腕伝送、Decimal式"),
            ("Fig. 7–8", "reproduce_noise_robustness.py", "脱分極厳密式、CHSH閾値"),
        ),
        (1.0, 2.65, 2.85),
        header_fill="E8EEF5",
    )
    add_figure(
        doc,
        "experiments/ding_jiang/results/fig3_v3/fig3_reproduction.png",
        "図1 Ding–Jiang v3 Figure 3再現。最大ギャップはp=0.5, beta=0で0.1035533906。",
        "Ding-Jiang HFT quantum-classical gap heatmap and cross-sections",
    )
    maximum = fig3["simulator_extrema"]["maximum_gap"]
    add_body(
        doc,
        f"Figure 3の101×101格子では最大Δω={maximum:.10f}です。beta=0.5でギャップ0、beta対称性、独立な16決定論戦略列挙がすべて許容差内で一致します。",
    )
    type2 = load_json(
        "experiments/ding_jiang/results/type_ii_memory_v3/type_ii_memory_summary.json"
    )["simulator"]
    add_table(
        doc,
        ("Type II量", "計算値", "論文丸め値"),
        (
            ("試行時間 t_a", f"{type2['attempt_time_us']:.6f} us", "約230 us"),
            ("成功確率 p_s", f"{type2['success_probability']:.10f}", "約0.0248"),
            ("メモリ当たり率", f"{type2['per_memory_rate_hz']:.6f} Hz", "約106 Hz"),
            ("100 Hzに必要なM", str(type2["minimum_memory_count_for_target"]), "1"),
        ),
        (2.0, 2.3, 2.2),
        header_fill="E8EEF5",
    )

    add_heading(doc, "7. Li et al.再現手順", 1)
    add_code(
        doc,
        "PYTHONPATH=src python3 -m experiments.li2026.reproduce_fig2\n"
        "PYTHONPATH=src python3 -m experiments.li2026.reproduce_fig3\n"
        "PYTHONPATH=src python3 experiments/li2026/reproduce_table3_50km.py\n"
        "PYTHONPATH=src python3 -m experiments.li2026.reproduce_fig7b",
    )
    add_bullets(
        doc,
        (
            "Figure 2: 独立Bernoulli入力、相関入力、対称/非対称効用、結合誤差と閾値。",
            "Figure 3: 厳密二項上側確率、最小n_req、T_envごとのR_req。",
            "Table III: 50 kmシステム値を下位パラメータから導出し、表示値との差も保持。",
            "Figure 7(b): 三者多数決パリティゲームとGHZ戦略の方程式レベル再現。",
        ),
    )
    add_figure(
        doc,
        "experiments/li2026/results/fig2_v1/fig2_reproduction.png",
        "図2 Li et al. Figure 2再現。入力相関、効用非対称、誤差がギャップへ与える影響。",
        "Li Figure 2 generalized LCTC gap reproduction",
    )

    add_heading(doc, "8. 運用上の量子優位判定", 1)
    add_heading(doc, "8.1 忠実度", 2)
    add_body(doc, "厳密結合式 ε=1-(1-4ε_s/3)(1-2ε_meas)^2 を用います。メモリを使う場合はτ_occに応じて状態誤差を更新した後、ε<ε_th=1-C/Qを判定します。")
    add_heading(doc, "8.2 有限統計", 2)
    add_body(doc, "勝敗効用には二項上側確率を用い、p<alphaとなる最小整数n_reqを探索します。一般得点にはLi Eq. 20の保守的境界を対数空間で評価します。探索上限到達と数学的に有限解がない場合は区別されます。")
    add_heading(doc, "8.3 判断遅延", 2)
    add_body(doc, "τ_decは基底選択/回転時間と測定時間の和です。等号は不合格であり、τ_dec<T_loc<T_commをそれぞれ独立に確認します。")
    case = table3["operational_cases"][0]
    add_table(
        doc,
        ("50 km / T_env=10 ms", "値", "状態"),
        (
            ("理論ギャップ", f"{case['ideal_gap']:.9f}", case["theoretical_advantage"]),
            ("誤差後ギャップ", f"{case['noisy_gap']:.9f}", case["fidelity_criterion"]),
            ("必要試行数", str(case["n_req"]), case["statistical_certification"]),
            ("R_HEG / R_req", f"{case['r_heg']:.1f} / {case['r_req']:.1f} s^-1", case["rate_criterion"]),
            ("τ_dec / T_loc", f"{case['tau_dec']*1e6:.3f} / {case['t_loc']*1e6:.1f} us", case["decision_criterion"]),
            ("総合", case["overall_operational_quantum_advantage"], case["overall_operational_quantum_advantage"]),
        ),
        (2.45, 2.55, 1.5),
        header_fill="E8EEF5",
    )

    add_heading(doc, "9. ハードウェア・M2イベント駆動モデル", 1)
    add_body(
        doc,
        "M0はメモリなし、M1はDing–Jiang型の一般メモリ、M2はLiのイベント準備型時間多重化です。M2では試行周期、リンク待ち、局所判断、リセット、有限メモリ数、チャネル数、寿命、デコヒーレンスを別々に保持します。",
    )
    system = table3["system_level_result"]
    add_table(
        doc,
        ("50 kmシステム量", "導出値", "注意"),
        (
            ("τ_dec", f"{system['tau_dec']*1e6:.3f} us", "100 ns + 870 ns"),
            ("τ_occ", f"{system['timing']['tau_occ']*1e6:.3f} us", "表示244 usとの差を記録"),
            ("p_ent", f"{system['entanglement_success_probability']:.8f}", "Eq. 56から導出"),
            ("R_HEG", f"{system['rate']['r_heg']:.3f} s^-1", "最終値は非ハードコード"),
            ("メモリ調整後ε", f"{system['memory_adjusted_combined_infidelity_upper_bound']:.9f}", "τ_mem=7.9 s"),
            ("必要メモリ深さ", str(system['timing']['minimum_memory_qubits']), "250個ではメモリ制限"),
        ),
        (2.0, 2.25, 2.25),
        header_fill="E8EEF5",
    )
    add_code(doc, "PYTHONPATH=src python3 experiments/li2026/cross_validate_m2_event_simulation.py")
    add_body(doc, "イベント駆動結果では、256 seed、26,368,000試行の平均Bell対生成率7863.867 s^-1、標準偏差301.451 s^-1、95%信頼区間[7826.764, 7900.970] s^-1となり、解析値7854.545 s^-1と整合しました。")

    add_heading(doc, "10. 結果の読み方", 1)
    add_heading(doc, "10.1 JSONサマリー", 2)
    add_bullets(
        doc,
        (
            "overall_statusは計算・内部検証ゲートの結果。",
            "paper_reproduction_statusは論文そのものに対するPASS/PARTIAL等。両者を混同しない。",
            "reference/versionで論文版を確認する。",
            "validationsにはactual、expected、誤差、許容差、provenanceを記録する。",
            "documented_discrepanciesは調整せず、そのまま研究上の差異として扱う。",
        ),
    )
    add_heading(doc, "10.2 CSVとPNG", 2)
    add_body(doc, "CSVは図の再描画・二次解析用であり、PNGは視覚比較用です。図だけをオラクルにせず、可能な限り式、列挙、公開数値を優先してください。Monte Carloではseed、標本数、平均、標準偏差、信頼区間を必ず併記します。")

    add_heading(doc, "11. 検証・回帰・拡張", 1)
    add_heading(doc, "11.1 最終検証", 2)
    add_code(
        doc,
        "python3 experiments/final_validation/run_phase15.py\n"
        "python3 -m pytest tests/scientific/test_phase15_final_validation_artifacts.py -q",
    )
    add_body(doc, "ランナーは全再生成を一時ディレクトリで行い、コミット済み成果物をオラクルとして比較します。結果ディレクトリが実行前後で同一ハッシュであることも確認します。")
    add_heading(doc, "11.2 新規モデル追加の最低条件", 2)
    add_bullets(
        doc,
        (
            "既存の層境界と設定駆動方式を維持する。",
            "古典基準を決定論戦略列挙など独立経路で確認する。",
            "解析モデルを先に実装し、イベント駆動モデルと比較する。",
            "Ding–Jiang回帰を毎回実行し、値が変化した場合は原因分類を記録する。",
            "新しい論文版は既存版の結果を上書きせず、別設定・別成果物として保存する。",
        ),
    )

    add_heading(doc, "12. トラブルシューティング", 1)
    add_table(
        doc,
        ("症状", "確認事項", "対応"),
        (
            ("import失敗", "実行場所、PYTHONPATH、editable install", "リポジトリルートでpip install -eを実行"),
            ("CVXPY solver差", "solver/version、許容差、bracket", "解析上下界を確認し、期待値を安易に更新しない"),
            ("n_req未決定", "ε>=ε_thか探索上限か", "no-finite-solutionとsearch-limitを区別"),
            ("rateだけFAIL", "T_env、N_channels、p_ent", "R_reqとR_HEGを別々に出力して律速を特定"),
            ("decisionだけFAIL", "τ_rot、τ_meas、T_loc", "通信時間ではなく局所判断時間を改善"),
            ("Table III表示差", "式入力と丸め順", "documented_discrepanciesを参照し、隠れ調整しない"),
            ("Monte Carlo差", "seed、試行数、CI、z値", "CIと解析値の両方で判定"),
        ),
        (1.3, 2.3, 2.9),
        header_fill="E8EEF5",
        font_size=8.0,
    )
    add_heading(doc, "付録A. 主要コマンド一覧", 1)
    add_code(
        doc,
        "python3 -m pytest\n"
        "PYTHONPATH=src python3 experiments/ding_jiang/reproduce_fig5_cross_sections.py\n"
        "PYTHONPATH=src python3 -m experiments.li2026.analyze_hft_waterfall\n"
        "PYTHONPATH=src python3 -m experiments.li2026.optimize_hardware_resources\n"
        "python3 experiments/performance/benchmark_phase14.py\n"
        "python3 experiments/final_validation/run_phase15.py",
    )
    add_heading(doc, "付録B. 科学的制約", 1)
    add_bullets(
        doc,
        (
            "一般の非二値ゲームに対する汎用量子最適化器は未実装。",
            "Figure 5の全面101×101明示戦略探索は最終ゲートで再実行せず、検証済み断面を再利用。",
            "Li Figure 2/3/7(b)は著者点データがないため方程式・形状レベルのPARTIAL。",
            "微視的TPI、測定、CAPSモデルは必要データ不足のためINSUFFICIENT_INFORMATION。",
            "ハードウェア最適化は設定済み有限格子上でのみ厳密で、連続大域最適性を主張しない。",
        ),
    )
    doc.save(MANUAL_PATH)


def build_report() -> None:
    phase15 = load_json(
        "experiments/final_validation/results/phase15_v1/final_validation_summary.json"
    )
    ding_fig3 = load_json("experiments/ding_jiang/results/fig3_v3/fig3_summary.json")
    type2 = load_json(
        "experiments/ding_jiang/results/type_ii_memory_v3/type_ii_memory_summary.json"
    )
    li_fig2 = load_json("experiments/li2026/results/fig2_v1/fig2_summary.json")
    li_fig3 = load_json("experiments/li2026/results/fig3_v1/fig3_summary.json")
    table3 = load_json("experiments/li2026/results/table3_50km_v1/table3_summary.json")
    event = load_json(
        "experiments/li2026/results/m2_event_cross_validation_v1/m2_event_summary.json"
    )
    waterfall = load_json(
        "experiments/li2026/results/hft_waterfall_v1/hft_waterfall_summary.json"
    )
    optimization = load_json(
        "experiments/li2026/results/hardware_optimization_v1/hardware_optimization_summary.json"
    )
    multiparty = load_json("experiments/li2026/results/fig7b_v1/fig7b_summary.json")

    doc = Document()
    configure_document(
        doc,
        title="量子ネットワーク・シミュレータ研究開発報告書",
        subject="Ding-Jiang to Li operational LCTC reproduction and extension report",
        justified=True,
    )
    add_cover(
        doc,
        kicker="RESEARCH & DEVELOPMENT REPORT",
        title="量子ネットワーク・シミュレータ",
        subtitle_lines=("Ding–Jiang再現からLi et al.運用LCTC拡張へ", "最終研究開発報告書"),
        audience="理論、有限統計、ネットワーク物理、HFT実現可能性の統合評価",
    )
    add_heading(doc, "要旨", 1)
    add_body(
        doc,
        "本研究は、Ding and Jiangの量子テレパシー型HFTモデルを再現し、Li et al.の運用基準へ連続的に拡張した。古典最適値は局所決定論戦略の完全列挙で独立検証し、量子値、物理誤差、有限統計、局所判断遅延、有限エンタングルメント供給、メモリ占有を分離した。これにより、正の量子–古典ギャップと、統計的に証明可能でハードウェア上実現可能な量子優位を明確に区別した。",
    )
    add_status_callout(
        doc,
        "最終検証: Phase 15 PASS。12/12再現ジョブ、592既存テスト、3最終成果物テストがPASS。",
    )
    add_table(
        doc,
        ("主要到達点", "結果"),
        (
            ("CHSH解析ベンチマーク", "C=0.75、Q=0.8535533906、Δω=0.1035533906"),
            ("Ding–Jiang v3", "HFT、損失、Type IIメモリ、ノイズを回帰化"),
            ("Li et al. v1", "一般化効用、相関入力、忠実度、有限統計、三基準を実装"),
            ("50 km M2", "R_HEG=7854.545 s^-1、τ_dec=0.970 us"),
            ("イベント照合", "256 seed、平均7863.867 s^-1、解析値との差0.1187%"),
            ("HFT運用分析", "8シナリオ中3 PASS、5制御FAIL"),
            ("資源最適化", "43,776候補、設定済み有限格子で7推奨設計"),
        ),
        (2.2, 4.3),
        header_fill="F4F6F9",
    )
    add_contents(
        doc,
        (
            "研究目的と文献版",
            "方法とソフトウェア構造",
            "Ding–Jiang再現",
            "Li一般化LCTCと忠実度",
            "有限統計と運用三基準",
            "50 km M2ハードウェア",
            "解析・イベント駆動クロスバリデーション",
            "HFT運用ウォーターフォール",
            "三者拡張と資源最適化",
            "最終検証、限界、結論",
        ),
    )

    add_heading(doc, "1. 研究目的と文献版", 1)
    add_body(doc, "研究目的は、理想的な非局所ゲーム優位を再計算することだけではない。市場観測後の通信を禁じる時間制約、状態・測定誤差、統計的有意性、有限の環境定常窓、供給率、局所測定時間を同一の判定系へ接続し、反証可能な運用可否を出力することである。")
    add_table(
        doc,
        ("文献", "固定版", "本研究での役割"),
        (
            ("Ding and Jiang", "arXiv:2407.21723v3", "理想化LCTC/HFT、損失、Type IIメモリ、脱分極"),
            ("Li et al.", "arXiv:2604.07451v1", "非対称効用、有限統計、三基準、M2/Yb、三者拡張"),
        ),
        (1.5, 1.85, 3.15),
        header_fill="F4F6F9",
    )
    add_body(doc, "Ding–Jiangはv2でHFT例、v3で量子メモリ計算が訂正されたため、版混在を避けた。Liはv1を固定し、式と表示値が一致しない項目を差異として保持した。")

    add_heading(doc, "2. 方法とソフトウェア構造", 1)
    add_body(doc, "一般非局所ゲームから物理ネットワークまでを六つの責務に分けた。各層は下位層の検証済み量を利用するが、Ding–Jiang回帰値を暗黙に変更しない。")
    add_table(
        doc,
        ("層", "入力", "出力・独立オラクル"),
        (
            ("非局所ゲーム", "P(x,y), u(a,b|x,y)", "期待効用、直接有限和"),
            ("XOR/CHSH", "2×2行列M", "C(M), Q(M)、解析CHSH、16戦略列挙"),
            ("Ding HFT", "p, beta, loss/noise", "理想・損失・ノイズギャップ"),
            ("Li LCTC", "beta1, beta2, 相関P", "一般化ギャップ、ε_th"),
            ("運用統計", "alpha, T_env", "p値、n_req、R_req"),
            ("ネットワーク", "距離、効率、M2資源", "R_HEG、τ_dec、デコヒーレンス"),
        ),
        (1.2, 2.15, 3.15),
        header_fill="F4F6F9",
    )
    add_body(doc, "古典側は16個の二者二値決定論戦略を完全列挙する。共有乱数はその凸包しか生成しないため、列挙最大値が許容される無通信古典戦略の最適値である。量子側は2×2 XOR解析、明示qubit戦略、NPA Q1+AB上界、密度行列トレースを用途別に組み合わせた。")

    add_heading(doc, "3. Ding–Jiang再現", 1)
    add_heading(doc, "3.1 HFT量子–古典ギャップ", 2)
    ding_max = ding_fig3["simulator_extrema"]["maximum_gap"]
    add_body(doc, f"Figure 3の101×101格子で最大ギャップは{ding_max:.10f}となり、CHSH解析値との差は5.55×10^-17であった。beta=0.5のゼロギャップ、beta反転対称性、Theorem 10断面、古典列挙がすべてPASSした。")
    add_figure(
        doc,
        "experiments/ding_jiang/results/fig3_v3/fig3_reproduction.png",
        "図1 Ding–Jiang v3 Figure 3の方程式ベース再現。",
        "Ding-Jiang HFT gap heatmap",
    )
    add_heading(doc, "3.2 損失とType IIメモリ", 2)
    add_body(doc, "p=0.3、beta=0.3の損失例では、閾値η*=0.9405975342、古典値0.79、量子値0.7918676876を得た。Figure 5断面では明示qubit戦略を下界、独立なNPA Q1+ABを上界として閾値を挟んだ。著者の全面格子および修正NPAコードが未公開のため、Figure 5はPARTIALである。")
    sim = type2["simulator"]
    add_table(
        doc,
        ("v3 Type II量", "シミュレータ", "論文", "状態"),
        (
            ("t_a", f"{sim['attempt_time_us']:.6f} us", "約230 us", "PASS"),
            ("p_s", f"{sim['success_probability']:.10f}", "約0.0248", "PASS"),
            ("r_e/M", f"{sim['per_memory_rate_hz']:.6f} Hz", "約106 Hz", "PASS"),
            ("100 Hzに必要なM", str(sim['minimum_memory_count_for_target']), "1", "PASS"),
        ),
        (1.35, 2.05, 1.7, 1.4),
        header_fill="F4F6F9",
    )
    add_body(doc, "このM1計算は伝搬待ちが支配する理想化モデルであり、占有、リセット、有限寿命、デコヒーレンスを含まない。それらはLiのM2モデルで別途導入した。")
    add_heading(doc, "3.3 脱分極ノイズ", 2)
    add_body(doc, "Ding Eq. 4.2の二値出力一様混合を厳密実装した。最大ノイズ閾値はCHSH点で0.2928932188であり、Figures 7–8の解析不変量と正領域の縮小を確認した。著者点データがないため図の状態はPARTIALを維持する。")

    add_heading(doc, "4. Li一般化LCTCと忠実度", 1)
    add_body(doc, "Liの拡張ではbeta1とbeta2を独立化し、任意の相関入力分布P(x,y)を受け取る。一般化行列の右下符号はEq. 24および一様CHSH極限と整合する負号を採用し、表示Eq. 25との差を文書化した。")
    panel_a = li_fig2["simulator_extrema"]["panel_a_maximum_gap"]
    panel_b = li_fig2["simulator_extrema"]["panel_b_maximum_gap"]
    add_table(
        doc,
        ("Figure 2領域", "最大ギャップ", "位置", "古典/量子値"),
        (
            ("独立Bernoulli", f"{panel_a['gap']:.10f}", "p=0.5, beta1=beta2=0", "0.75 / 0.853553"),
            ("相関入力", f"{panel_b['gap']:.10f}", "P11=0.4, beta1=beta2=0", "0.8 / 0.867423"),
            ("CHSH ε_th", "0.2928932188", "1-1/sqrt(2)", "厳密一致"),
        ),
        (1.45, 1.65, 2.1, 1.3),
        header_fill="F4F6F9",
    )
    add_figure(
        doc,
        "experiments/li2026/results/fig2_v1/fig2_reproduction.png",
        "図2 Li et al. Figure 2。独立/相関入力と結合インフィデリティ依存。",
        "Li Figure 2 reproduction panels",
    )
    add_body(doc, "状態誤差と測定誤差は ε=1-(1-4ε_s/3)(1-2ε_meas)^2 で結合する。ε_s=0.04、ε_meas=0.002ではε=0.06089152であり、小誤差近似に置き換えていない。")

    add_heading(doc, "5. 有限統計と運用三基準", 1)
    add_body(doc, "勝敗効用では古典最適値を帰無仮説の勝率として、観測勝数以上の厳密二項上側確率を計算する。p<alphaとなる最小nをn_reqとし、R_req=n_req/T_envを要求する。小nは60桁Decimal直接和で独立照合した。")
    points = li_fig3["reference_line_points"]
    rows = []
    for point in points:
        rows.append(
            (
                f"{point['alpha']:.3g}",
                f"{point['t_env_seconds']:.3g}",
                str(point["required_trials"]),
                f"{point['required_rate_hz']:.1f}",
                "PASS" if point["reference_hardware_rate_pass"] else "FAIL",
            )
        )
    add_table(
        doc,
        ("alpha", "T_env [s]", "n_req", "R_req [s^-1]", "7.9 kHz"),
        rows,
        (0.8, 1.1, 1.0, 1.8, 1.8),
        header_fill="F4F6F9",
    )
    add_figure(
        doc,
        "experiments/li2026/results/fig3_v1/fig3_reproduction.png",
        "図3 Li et al. Figure 3。εが忠実度閾値へ近づくと必要率が急増する。",
        "Li Figure 3 finite-statistics required rates",
    )
    add_body(doc, "三基準は、(A) ε<ε_th、(B) R_HEG>R_req、(C) τ_dec<T_locである。さらに無通信LCTCとしてT_loc<T_commを要求する。いずれも厳密不等号で評価し、統計的認証と総合運用状態を別フィールドで返す。")

    add_heading(doc, "6. 50 km M2ハードウェア", 1)
    system = table3["system_level_result"]
    memory = table3["memory_fidelity_result"]
    add_body(doc, "M2はイベント準備型時間多重化を、有限メモリ占有、リンク遅延、再利用時間、メモリ数、チャネル数、寿命とともにモデル化する。50 km Ybシステム例は、最終R_HEGをハードコードせず、発光・TPI・光学・検出・リンク損失から導出した。")
    add_table(
        doc,
        ("量", "導出値", "論文表示", "再現状態"),
        (
            ("τ_dec", f"{system['tau_dec']*1e6:.3f} us", "0.97 us", "PASS"),
            ("τ_occ", f"{system['timing']['tau_occ']*1e6:.3f} us", "244 us", "PARTIAL"),
            ("p_ent", f"{system['entanglement_success_probability']:.8f}", "0.0077", "PARTIAL"),
            ("R_HEG", f"{system['rate']['r_heg']:.3f} s^-1", "7.9×10^3 s^-1", "PASS"),
            ("p_false", f"{system['false_positive_fraction']*100:.6f}%", "0.12%", "PARTIAL"),
            ("メモリ調整後ε", f"{memory['epsilon']:.9f}", "<6.1%", "PASS"),
        ),
        (1.25, 1.85, 1.75, 1.65),
        header_fill="F4F6F9",
    )
    add_body(doc, "Table IIIの11指標中7指標が表示丸め区間内でPASSし、4指標がPARTIALとなった。R0はp_e^2/2=0.245を先に0.25へ丸めると表示値へ近づく。p_ent、τ_occ、p_falseも式の非丸め計算と表示値が半単位区間で一致しない。隠れパラメータによる調整は行っていない。")
    op10 = table3["operational_cases"][0]
    op100 = table3["operational_cases"][1]
    add_table(
        doc,
        ("T_env", "n_req", "R_req", "R_HEG", "総合"),
        (
            ("10 ms", str(op10['n_req']), f"{op10['r_req']:.0f} s^-1", f"{op10['r_heg']:.1f} s^-1", op10['overall_operational_quantum_advantage']),
            ("100 ms", str(op100['n_req']), f"{op100['r_req']:.0f} s^-1", f"{op100['r_heg']:.1f} s^-1", op100['overall_operational_quantum_advantage']),
        ),
        (1.0, 1.0, 1.45, 1.7, 1.35),
        header_fill="F4F6F9",
    )

    add_heading(doc, "7. 解析・イベント駆動クロスバリデーション", 1)
    analytical = event["analytical_oracle"]
    monte = event["monte_carlo"]
    rate_stats = monte["bell_pair_rate_statistics"]
    add_body(doc, "解析式と独立な離散イベントスケジューラを比較した。固定タイミングと独立Bernoulli heraldの仮定下で、各seedにつき有限観測窓をシミュレートし、再現可能なseed列を保存した。")
    add_table(
        doc,
        ("統計量", "値"),
        (
            ("seed数", str(monte["replicates"])),
            ("総試行数", f"{monte['total_heralded_trials']:,}"),
            ("解析R_HEG", f"{analytical['bell_pair_rate_hz']:.3f} s^-1"),
            ("イベント平均", f"{rate_stats['mean']:.3f} s^-1"),
            ("標本標準偏差", f"{rate_stats['sample_std']:.3f} s^-1"),
            ("95%信頼区間", f"[{rate_stats['confidence_interval_lower']:.3f}, {rate_stats['confidence_interval_upper']:.3f}] s^-1"),
            ("解析値との差", "9.322 s^-1 (0.1187%)"),
            ("標準化残差", f"{monte['standardized_binomial_residual']:.3f}"),
        ),
        (2.35, 4.15),
        header_fill="F4F6F9",
    )
    add_body(doc, "解析値は95%信頼区間内にあり、平均差は事前許容差2%を十分下回った。ジッタ、時間相関故障、キュー競合は本検証の対象外である。")

    add_heading(doc, "8. HFT運用ウォーターフォール", 1)
    add_body(doc, "理想Ding利得、Li一般化、誤差後利得、有限統計、HEG率、判断遅延、無通信条件を順に適用した。8シナリオ中3件が総合PASS、5件が意図した律速でFAILした。")
    add_figure(
        doc,
        "experiments/li2026/results/hft_waterfall_v1/hft_waterfall.png",
        "図4 HFT理想利得から運用判定までのウォーターフォール。",
        "HFT operational advantage waterfall",
    )
    waterfall_rows = []
    for result in waterfall["results"]:
        waterfall_rows.append(
            (
                result["scenario_id"].replace("_", " "),
                f"{result['noisy_gap']:.6f}",
                "-" if result["n_req"] is None else str(result["n_req"]),
                result["dominant_bottleneck"].replace("_", " "),
                result["overall_operational_quantum_advantage"],
            )
        )
    add_table(
        doc,
        ("シナリオ", "誤差後Δω", "n_req", "主要律速", "総合"),
        waterfall_rows,
        (2.15, 1.0, 1.0, 1.55, 0.8),
        header_fill="F4F6F9",
        font_size=7.6,
    )
    add_body(doc, "Ding代表p=0.3、beta=0.3は、誤差後ギャップ0.00329536、n_req=66,133である。T_env=10 sならR_req=6,613.3 s^-1で50 km供給率を下回りPASSするが、1 sでは66,133 s^-1となりrate FAILへ転じる。この例は理論利得が正でも環境定常窓が短ければ運用上無意味になり得ることを示す。")

    add_heading(doc, "9. 三者拡張と資源最適化", 1)
    add_heading(doc, "9.1 三者GHZ/XOR", 2)
    mpmax = multiparty["simulator_extrema"]["maximum"]
    mpcase = multiparty["representative_operational_case"]["result"]
    add_body(doc, f"三者多数決パリティゲームでは64古典決定論戦略を列挙し、GHZ三位相最適化と独立な定常多項式経路を照合した。101×101格子の最大ギャップは{mpmax['gap']:.10f}、表示閾値を超える正ギャップ点は3,474である。")
    add_figure(
        doc,
        "experiments/li2026/results/fig7b_v1/fig7b_reproduction.png",
        "図5 Li et al. Figure 7(b)の三者ギャップ再現。",
        "Three-party GHZ XOR gap reproduction",
    )
    add_body(doc, f"代表感度ケースは結合誤差{mpcase['combined_infidelity']:.7f}、n_req={mpcase['n_req']}、供給率1 MHzで総合PASSとなる。ただし1 MHzは供給値であり、Appendix CのCAPS微視的再現ではない。")
    add_heading(doc, "9.2 有限格子ハードウェア最適化", 2)
    search = optimization["search_space"]
    envelope = optimization["distance_envelope"]
    add_body(doc, f"9ケース、各{search['candidate_count_per_case']:,}候補、合計43,776候補を全探索し、独立Pareto判定を行った。コストはTable III基準から探索境界までの正規化改善量であり、物理的費用ではない。")
    add_figure(
        doc,
        "experiments/li2026/results/hardware_optimization_v1/hardware_optimization.png",
        "図6 有限格子の最小改善量と距離別実現可能性。",
        "Hardware resource optimization and distance envelope",
    )
    add_body(doc, f"Ding 1 s rate-stressケースでは50 kmで最小9チャネル、探索範囲内で最大{envelope['maximum_configured_feasible_distance_km']:.0f} kmまで実現可能、{envelope['first_configured_infeasible_distance_km']:.0f} kmで不可能となった。これは設定済み有限格子上の結果であり、連続大域最適性を主張しない。")

    add_heading(doc, "10. 最終検証、限界、結論", 1)
    add_heading(doc, "10.1 Phase 15", 2)
    matrix = phase15["validations"]["reproduction_matrix"]
    add_table(
        doc,
        ("監査", "結果"),
        (
            ("再現/拡張ジョブ", "12/12 PASS"),
            ("既存テスト", phase15["validations"]["test_suite"]["summary"]),
            ("成果物不変性", phase15["validations"]["committed_artifacts_unchanged"]["status"]),
            ("成果物ファイル", str(phase15["committed_artifact_manifest"]["file_count"])),
            ("再現行列", f"{matrix['row_count']}行、語彙/形状 {matrix['status']}"),
            ("行列状態内訳", ", ".join(f"{key}={value}" for key, value in matrix["status_counts"].items())),
        ),
        (2.15, 4.35),
        header_fill="F4F6F9",
    )
    add_heading(doc, "10.2 科学的限界", 2)
    add_bullets(
        doc,
        (
            "著者の点データがないFigure 2、3、7(b)およびDing Figures 3、5、7、8は方程式・不変量・視覚レベルのPARTIAL。",
            "Table IIIは最終率を再現するが、R0、p_ent、τ_occ、p_falseの表示丸め差をPARTIALとして保持。",
            "微視的TPI、測定、CAPSはパルス情報または著者コード不足によりINSUFFICIENT_INFORMATION。",
            "一般非二値ゲーム向け汎用Bell演算子最適化はNOT_IMPLEMENTED。",
            "市場入力分布と効用は感度シナリオであり、実市場データによる較正・収益主張ではない。",
            "イベントモデルは固定タイミングと独立heraldを仮定し、故障相関や運用制御面を含まない。",
        ),
    )
    add_heading(doc, "10.3 結論", 2)
    add_body(doc, "Ding–Jiangの理想HFT優位は独立古典列挙と解析量子値により再現され、Liの一般化、忠実度、有限統計、判断遅延、M2供給へ後方互換的に接続された。50 km代表ケースでは、式から導出した約7.85 kHzのHEG率と0.97 usの判断遅延により、設定した10 msおよび100 msのCHSH定常窓で総合PASSを得た。一方、Ding代表HFTではギャップが小さく、T_envを10 sから1 sへ短縮すると率制約でFAILする。支配的律速はゲーム・市場条件ごとに変化し、理論優位だけでは運用可否を判断できない。")
    add_body(doc, "今後の研究優先順位は、著者データによるPARTIAL図の点wise照合、微視的光子/測定/CAPSモデルの情報取得、相関故障を含むイベントモデル、実市場データでのP(x,y)と効用較正、連続多目的最適化である。これらは現在のPASS境界を変更せず、新しい版・設定・オラクルとして追加する。")

    add_heading(doc, "参考文献", 1)
    add_body(doc, "[1] D. Ding and L. Jiang, “Coordinating Decisions via Quantum Telepathy,” arXiv:2407.21723v3, revised August 26, 2025.")
    add_body(doc, "[2] C. Li, S. Kikura, A. Goban, H. Yamasaki, and S. Sunami, “Operational criteria for quantum advantage in latency-constrained nonlocal games,” arXiv:2604.07451v1, April 8, 2026.")
    add_body(doc, "[3] J. F. Clauser, M. A. Horne, A. Shimony, and R. A. Holt, “Proposed experiment to test local hidden-variable theories,” Physical Review Letters 23, 880–884 (1969).")
    add_body(doc, "[4] Reproduction matrices, model maps, configurations, oracles, and final-validation records contained in this repository, retrieval/validation dates 2026-08-31 through 2026-09-03.")
    doc.save(REPORT_PATH)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_manual()
    build_report()
    documents = [
        {
            "path": MANUAL_PATH.relative_to(ROOT).as_posix(),
            "bytes": MANUAL_PATH.stat().st_size,
            "sha256": sha256(MANUAL_PATH),
            "preset": "compact_reference_guide",
            "header_pattern": "editorial_cover",
        },
        {
            "path": REPORT_PATH.relative_to(ROOT).as_posix(),
            "bytes": REPORT_PATH.stat().st_size,
            "sha256": sha256(REPORT_PATH),
            "preset": "narrative_proposal",
            "header_pattern": "editorial_cover",
        },
    ]
    if PAPER_PATH.exists():
        documents.append(
            {
                "path": PAPER_PATH.relative_to(ROOT).as_posix(),
                "bytes": PAPER_PATH.stat().st_size,
                "sha256": sha256(PAPER_PATH),
                "preset": "latex_research_article",
                "header_pattern": "article_maketitle",
            }
        )
    manifest = {
        "schema_version": 1,
        "generated_on": "2026-09-03",
        "source_validation": "phase15_v1",
        "documents": documents,
    }
    manifest_path = OUTPUT_DIR / "phase16_document_manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    for item in manifest["documents"]:
        print(f"Created {item['path']} ({item['bytes']} bytes)")


if __name__ == "__main__":
    main()
