"""Structural regression tests for the three Phase 16 narrative deliverables."""

from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
DELIVERABLES = ROOT / "deliverables" / "phase16"
MANUAL = DELIVERABLES / "quantum_telepathy_simulator_manual_ja.docx"
REPORT = DELIVERABLES / "quantum_telepathy_research_report_ja.docx"
PAPER = DELIVERABLES / "operational_lctc_quantum_advantage.tex"
BIBLIOGRAPHY = DELIVERABLES / "references.bib"


def _document_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + table_cells)


def _docx_archive_facts(path: Path) -> dict[str, object]:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        document_xml = archive.read("word/document.xml").decode("utf-8")
    return {
        "has_header": any(name.startswith("word/header") for name in names),
        "has_footer": any(name.startswith("word/footer") for name in names),
        "media_count": sum(name.startswith("word/media/") for name in names),
        "alt_text_count": document_xml.count("descr="),
    }


def test_phase16_contains_exactly_three_narrative_documents() -> None:
    narrative_files = sorted(
        path.name
        for path in DELIVERABLES.iterdir()
        if path.suffix.lower() in {".docx", ".tex"}
    )
    assert narrative_files == sorted(
        [MANUAL.name, REPORT.name, PAPER.name]
    )


def test_japanese_word_documents_are_structurally_complete() -> None:
    phase15 = json.loads(
        (
            ROOT
            / "experiments/final_validation/results/phase15_v1/final_validation_summary.json"
        ).read_text(encoding="utf-8")
    )
    assert phase15["overall_status"] == "PASS"

    expectations = (
        (
            MANUAL,
            "量子テレパシー・シミュレータ",
            2,
            ("既存テスト 592 件 PASS", "12. トラブルシューティング"),
        ),
        (
            REPORT,
            "量子ネットワーク・シミュレータ",
            6,
            ("最終検証: Phase 15 PASS", "10.3 結論"),
        ),
    )
    forbidden = ("TODO", "PLACEHOLDER", "Lorem ipsum")
    for path, title, minimum_media, required_text in expectations:
        assert path.stat().st_size > 100_000
        document = Document(path)
        text = _document_text(path)
        facts = _docx_archive_facts(path)
        assert title in text
        assert all(value in text for value in required_text)
        assert not any(value in text for value in forbidden)
        assert sum(p.style.name == "Heading 1" for p in document.paragraphs) >= 10
        assert len(document.tables) >= 5
        assert facts["has_header"] is True
        assert facts["has_footer"] is True
        assert facts["media_count"] >= minimum_media
        assert facts["alt_text_count"] >= minimum_media


def test_english_latex_paper_has_resolved_local_inputs() -> None:
    source = PAPER.read_text(encoding="utf-8")
    bibliography = BIBLIOGRAPHY.read_text(encoding="utf-8")
    figures = re.findall(r"\\includegraphics(?:\[[^]]*\])?\{([^}]+)\}", source)
    citation_keys = {
        key.strip()
        for group in re.findall(r"\\cite\{([^}]+)\}", source)
        for key in group.split(",")
    }
    bibliography_keys = set(re.findall(r"@\w+\{([^,]+),", bibliography))

    assert "All 12 jobs pass" in source
    assert "592 pre-existing tests" in source
    assert "\\bibliography{references}" in source
    assert len(figures) == 6
    assert all((DELIVERABLES / "figures" / figure).is_file() for figure in figures)
    assert citation_keys <= bibliography_keys
    assert {"dingjiang2025", "li2026", "chsh1969", "npa2007"} <= bibliography_keys
    assert not any(token in source for token in ("TODO", "PLACEHOLDER", "Lorem ipsum"))
